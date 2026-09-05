"""
Word (.docx) report exporter using python-docx with structured tables,
typography, audit logs, and compliance disclaimer.
"""
import os
from typing import Dict, Any

def export_to_word(calc_data: Dict[str, Any], output_path: str) -> str:
    """Exports calculation results to a Microsoft Word (.docx) document."""
    try:
        import docx
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT

        doc = docx.Document()

        # Document margins
        for section in doc.sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)

        # Title
        title_p = doc.add_paragraph()
        title_run = title_p.add_run("INDIAN INCOME-TAX COMPUTATION REPORT")
        title_run.bold = True
        title_run.font.size = Pt(20)
        title_run.font.color.rgb = RGBColor(30, 58, 138)
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        sub_p = doc.add_paragraph()
        sub_run = sub_p.add_run("Comparative Tax Analysis & Statutory Estimation Workstation")
        sub_run.font.size = Pt(12)
        sub_run.font.color.rgb = RGBColor(100, 116, 139)
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # spacer

        # Taxpayer Details Table
        doc.add_heading("1. Taxpayer Profile & Regime Parameters", level=2)
        p_table = doc.add_table(rows=6, cols=2)
        p_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        p_data = [
            ("Taxpayer Name", calc_data.get("taxpayer_name", "N/A")),
            ("Permanent Account Number (PAN)", calc_data.get("pan_masked", "N/A")),
            ("Governing Tax Statute", calc_data.get("act_name", "Income-tax Act, 1961")),
            ("Tax Regime", calc_data.get("regime", "New Regime")),
            ("Financial Year (FY)", calc_data.get("financial_year", "2024-25")),
            ("Assessment Year (AY)", calc_data.get("assessment_year", "2025-26")),
        ]
        for idx, (label, val) in enumerate(p_data):
            p_table.rows[idx].cells[0].paragraphs[0].add_run(label).bold = True
            p_table.rows[idx].cells[1].paragraphs[0].add_run(str(val))

        doc.add_paragraph()

        # Core Computation Table
        doc.add_heading("2. Tax Liability Computation Breakdown", level=2)
        summary = calc_data.get("summary", {})
        comp_data = [
            ("Gross Total Income (GTI)", summary.get("gross_total_income", "0")),
            ("Less: Eligible Deductions", summary.get("total_deductions", "0")),
            ("Net Taxable Total Income", summary.get("taxable_total_income", "0")),
            ("Normal Slab Tax", summary.get("normal_slab_tax", "0")),
            ("Special Rate Tax (STCG 111A / LTCG 112A)", summary.get("special_rate_tax", "0")),
            ("Total Tax Payable before Rebate", summary.get("tax_before_rebate", "0")),
            ("Less: Rebate u/s 87A", summary.get("rebate_amount", "0")),
            ("Tax after Rebate", summary.get("tax_after_rebate", "0")),
            ("Add: Surcharge", summary.get("surcharge_amount", "0")),
            ("Add: Health & Education Cess (4%)", summary.get("cess_amount", "0")),
            ("Total Tax Liability (Rounded u/s 288B)", summary.get("total_tax_liability", "0")),
            ("Less: Prepaid Taxes & Credits (TDS/TCS/Advance Tax)", summary.get("total_prepaid_tax", "0")),
            ("Balance Tax Payable / (Refund)", summary.get("balance_tax_payable", "0")),
        ]

        c_table = doc.add_table(rows=len(comp_data) + 1, cols=2)
        c_table.rows[0].cells[0].paragraphs[0].add_run("Computation Head").bold = True
        c_table.rows[0].cells[1].paragraphs[0].add_run("Amount (INR)").bold = True

        for idx, (head, amt) in enumerate(comp_data):
            row_cells = c_table.rows[idx + 1].cells
            r_run = row_cells[0].paragraphs[0].add_run(head)
            if "Total" in head or "Balance" in head:
                r_run.bold = True
            val_run = row_cells[1].paragraphs[0].add_run(amt)
            if "Total" in head or "Balance" in head:
                val_run.bold = True

        doc.add_paragraph()

        # Audit Trail & Rules Applied
        doc.add_heading("3. Verified Tax Rules Applied & Computational Audit Trail", level=2)
        rules = calc_data.get("rules_applied", [])
        if rules:
            doc.add_paragraph("Statutory Rules Applied:", style='List Bullet')
            for r in rules:
                doc.add_paragraph(r, style='List Bullet')

        audit_steps = calc_data.get("audit_trail", [])
        if audit_steps:
            doc.add_paragraph()
            doc.add_paragraph("Calculation Trail Steps:")
            for s in audit_steps:
                doc.add_paragraph(s, style='List Number')

        doc.add_paragraph()

        # Mandatory Legal Disclaimer
        doc.add_heading("4. Legal and Compliance Safeguard Notice", level=2)
        disc_p = doc.add_paragraph()
        disc_run = disc_p.add_run(calc_data.get("disclaimer", ""))
        disc_run.italic = True
        disc_run.font.color.rgb = RGBColor(185, 28, 28)

        doc.save(output_path)
        return output_path

    except ImportError:
        # Fallback to plain formatted text report
        txt_path = output_path.replace(".docx", ".txt")
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write("=" * 60 + "\n")
            f.write("INDIAN INCOME-TAX COMPUTATION REPORT\n")
            f.write(f"Generated On: {calc_data.get('timestamp', '')}\n")
            f.write("=" * 60 + "\n\n")
            f.write(f"Taxpayer Name: {calc_data.get('taxpayer_name', '')}\n")
            f.write(f"PAN: {calc_data.get('pan_masked', '')}\n")
            f.write(f"Act: {calc_data.get('act_name', '')}\n")
            f.write(f"Regime: {calc_data.get('regime', '')}\n")
            f.write(f"FY: {calc_data.get('financial_year', '')} | AY: {calc_data.get('assessment_year', '')}\n\n")
            f.write("COMPUTATION SUMMARY:\n")
            for k, v in calc_data.get("summary", {}).items():
                f.write(f"  {k}: Rs. {v}\n")
            f.write("\nLEGAL DISCLAIMER:\n")
            f.write(calc_data.get("disclaimer", "") + "\n")
        return txt_path
